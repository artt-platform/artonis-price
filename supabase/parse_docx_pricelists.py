"""Parse DOCX files for the 2 known missing exhibitions:
   #8 Bùi Tiến Tuấn - Một Hành Trình (2 docx versions)
   #21 Ngô Thanh Hùng - Hơi Thở Của Đất
Uses Claude Sonnet 4.6 over the docx text content."""
import os, sys, json, re, unicodedata
from pathlib import Path
import requests
from docx import Document  # python-docx

sys.path.insert(0,'/Users/trietnguyen/Documents/Company/Artonis/App/ArtonisArtistPrice')
os.chdir('/Users/trietnguyen/Documents/Company/Artonis/App/ArtonisArtistPrice')
from artonis_price_mvp import normalize_key

ENV={}
for line in open('.env.local'):
    line=line.strip()
    if line and not line.startswith('#') and '=' in line:
        k,v=line.split('=',1); ENV[k]=v
URL=ENV['SUPABASE_URL']; KEY=ENV['SUPABASE_SERVICE_ROLE_KEY']
ANTH_KEY=ENV['ANTHROPIC_API_KEY']
H={'apikey':KEY,'Authorization':f'Bearer {KEY}','Content-Type':'application/json','Prefer':'return=minimal'}

import anthropic
client = anthropic.Anthropic(api_key=ANTH_KEY)
DRY = os.environ.get('DRY','0')=='1'

arts = requests.get(f'{URL}/rest/v1/artists?select=id,name,normalized_name', headers={'apikey':KEY}).json()
artist_by_norm = {a['normalized_name']: a['id'] for a in arts}

ARTWORK_SCHEMA = {
    "type":"object","properties":{
        "artworks":{"type":"array","items":{
            "type":"object","properties":{
                "artist_name":{"type":"string"},
                "artwork_title":{"type":"string"},
                "dimensions":{"type":"string"},
                "medium":{"type":"string"},
                "year":{"type":"string"},
                "price_amount":{"type":"number"},
                "currency":{"type":"string"},
                "status":{"type":"string"},
            },
            "required":["artist_name","artwork_title","price_amount","currency"],
            "additionalProperties":False,
        }}
    },
    "required":["artworks"],"additionalProperties":False,
}

SYSTEM_PROMPT = """Extract artwork prices from Vietnamese exhibition catalog text (often tabular).

For each priced artwork return:
- artist_name (Vietnamese diacritics)
- artwork_title
- dimensions (e.g. "60x80 cm")
- medium (chất liệu)
- year
- price_amount (number, strip commas/dots/đ)
- currency (VND / USD / EUR)
- status (sold / available / empty)

Skip items without explicit prices. Currency: "350.000.000" → 350000000 VND; "$3,500" → 3500 USD.
"""

def docx_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip(): parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = ' | '.join(c.text.strip() for c in row.cells)
            if row_text.strip(): parts.append(row_text)
    return '\n'.join(parts)

CASES = [
    {
        'exh_id': 8, 'artist_name': 'Bùi Tiến Tuấn',
        'paths': [
            '/tmp/artonis_drive/Bùi Tiến Tuấn - Một Hành Trình - 250222 - Sann/Bùi Tiến Tuấn (Q1) - Một Hành Trình - 20250222 - Sann.docx',
            '/tmp/artonis_drive/Bùi Tiến Tuấn - Một Hành Trình - 250222 - Sann/Bùi Tiến Tuấn (Q2) - Một Hành Trình - 22022025 - Sann.docx',
        ],
        'title': 'Một Hành Trình',
    },
    {
        'exh_id': 21, 'artist_name': 'Ngô Thanh Hùng',
        'paths': ['/tmp/artonis_drive/Ngô Thanh Hùng - Hơi Thở Của Đất - 251130 - Taa/Ngô Thanh Hùng - Hơi Thở Của Đất - 20251130 - Taa.docx'],
        'title': 'Hơi Thở Của Đất',
    },
]

all_rows = []
total_in=0; total_out=0
for case in CASES:
    aid = artist_by_norm.get(normalize_key(case['artist_name']))
    if not aid:
        print(f"  no artist id for {case['artist_name']}"); continue
    text_parts = []
    for p in case['paths']:
        if not os.path.exists(p):
            print(f"  missing: {p}"); continue
        try:
            text_parts.append(f"=== {os.path.basename(p)} ===\n{docx_text(p)}")
        except Exception as e:
            print(f"  parse err {p}: {e}"); continue
    if not text_parts:
        continue
    combined = '\n\n'.join(text_parts)
    print(f"\n→ #{case['exh_id']} {case['artist_name']} | {len(combined)} chars")
    print(f"   sample: {combined[:200]}")
    if DRY: continue
    resp = client.messages.create(
        model="claude-sonnet-4-6", max_tokens=8000,
        output_config={"effort":"low","format":{"type":"json_schema","schema":ARTWORK_SCHEMA}},
        system=[{"type":"text","text":SYSTEM_PROMPT,"cache_control":{"type":"ephemeral"}}],
        messages=[{"role":"user","content":f"Exhibition: {case['title']}\nArtist: {case['artist_name']}\n\n{combined[:60000]}"}],
    )
    total_in += resp.usage.input_tokens
    total_out += resp.usage.output_tokens
    out = next(b.text for b in resp.content if b.type=='text')
    artworks = json.loads(out)['artworks']
    print(f"   {len(artworks)} priced (in={resp.usage.input_tokens}, out={resp.usage.output_tokens})")
    for a in artworks:
        all_rows.append({
            'artist_id': aid, 'exhibition_id': case['exh_id'],
            'artwork_title': (a.get('artwork_title') or '')[:300],
            'medium': a.get('medium') or None,
            'dimensions': a.get('dimensions') or '',
            'year': a.get('year') or None,
            'price_amount': a.get('price_amount'),
            'currency': a.get('currency') or 'VND',
            'status': a.get('status') or '',
            'confidence': 0.85,
        })

print(f"\nrows={len(all_rows)} tokens in={total_in} out={total_out} cost~${total_in/1e6*3+total_out/1e6*15:.3f}")
if all_rows and not DRY:
    rsp = requests.post(f'{URL}/rest/v1/price_observations', headers=H, json=all_rows, timeout=60)
    print(f'insert HTTP {rsp.status_code}')
    if rsp.status_code not in (200,201,204): print(rsp.text[:300])
